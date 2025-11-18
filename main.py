import os
from datetime import datetime, timedelta, timezone
from typing import List, Optional
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.staticfiles import StaticFiles
from jose import JWTError, jwt
from passlib.context import CryptContext
from pydantic import BaseModel
from bson import ObjectId
import qrcode
from io import BytesIO
from PIL import Image
import base64

from database import db, create_document, get_documents
from schemas import User, Drink, TheorySection, Certificate, ShareLink

SECRET_KEY = os.getenv("SECRET_KEY", "dev-secret-change-me")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24

app = FastAPI(title="Bartender Academy API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Static files for uploads
UPLOAD_DIR = os.getenv("UPLOAD_DIR", "uploads")
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR, exist_ok=True)
app.mount("/static", StaticFiles(directory=UPLOAD_DIR), name="static")

# Auth setup
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="/auth/token")


class Token(BaseModel):
    access_token: str
    token_type: str = "bearer"


class TokenData(BaseModel):
    email: Optional[str] = None


# Utilities

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return pwd_context.verify(plain_password, hashed_password)


def get_password_hash(password: str) -> str:
    return pwd_context.hash(password)


def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)


def get_user_by_email(email: str) -> Optional[dict]:
    if db is None:
        return None
    return db["user"].find_one({"email": email})


def get_current_user(token: str = Depends(oauth2_scheme)) -> dict:
    credentials_exception = HTTPException(status_code=401, detail="Could not validate credentials")
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    user = get_user_by_email(email)
    if user is None:
        raise credentials_exception
    return user


def get_current_admin(user: dict = Depends(get_current_user)) -> dict:
    if not user.get("is_admin", False):
        raise HTTPException(status_code=403, detail="Admin privileges required")
    return user


@app.get("/")
def root():
    return {"message": "Bartender Academy API running"}


@app.get("/test")
def test_database():
    response = {
        "backend": "✅ Running",
        "database": "❌ Not Available",
        "database_url": None,
        "database_name": None,
        "connection_status": "Not Connected",
        "collections": []
    }
    try:
        if db is not None:
            response["database"] = "✅ Available"
            response["database_name"] = db.name
            response["connection_status"] = "Connected"
            collections = db.list_collection_names()
            response["collections"] = collections
            response["database"] = "✅ Connected & Working"
        else:
            response["database"] = "⚠️  Available but not initialized"
    except Exception as e:
        response["database"] = f"❌ Error: {str(e)[:80]}"
    response["database_url"] = "✅ Set" if os.getenv("DATABASE_URL") else "❌ Not Set"
    response["database_name"] = "✅ Set" if os.getenv("DATABASE_NAME") else "❌ Not Set"
    return response


# Auth endpoints
@app.post("/auth/register", response_model=Token)
def register(name: str = Form(...), email: str = Form(...), password: str = Form(...)):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    existing = db["user"].find_one({"email": email})
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")
    password_hash = get_password_hash(password)
    user = User(name=name, email=email, password_hash=password_hash, is_admin=True)
    user_id = create_document("user", user)
    token = create_access_token({"sub": email})
    return Token(access_token=token)


@app.post("/auth/token", response_model=Token)
def login(form_data: OAuth2PasswordRequestForm = Depends()):
    user = get_user_by_email(form_data.username)
    if not user or not verify_password(form_data.password, user.get("password_hash", "")):
        raise HTTPException(status_code=401, detail="Incorrect email or password")
    token = create_access_token({"sub": user["email"]})
    return Token(access_token=token)


# Drinks endpoints
@app.get("/drinks")
def list_drinks(q: Optional[str] = None, method: Optional[str] = None, base_spirit: Optional[str] = None, category: Optional[str] = None, glassware: Optional[str] = None):
    if db is None:
        return []
    query = {}
    if q:
        query["$or"] = [
            {"name_it": {"$regex": q, "$options": "i"}},
            {"name_en": {"$regex": q, "$options": "i"}},
        ]
    if method:
        query["method"] = method
    if base_spirit:
        query["base_spirit"] = base_spirit
    if category:
        query["category"] = category
    if glassware:
        query["glassware"] = glassware
    items = list(db["drink"].find(query).sort("name_en"))
    for it in items:
        it["_id"] = str(it["_id"])  # jsonify
    return items


@app.get("/drinks/{drink_id}")
def get_drink(drink_id: str):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    try:
        doc = db["drink"].find_one({"_id": ObjectId(drink_id)})
    except Exception:
        raise HTTPException(status_code=404, detail="Drink not found")
    if not doc:
        raise HTTPException(status_code=404, detail="Drink not found")
    doc["_id"] = str(doc["_id"])  # jsonify
    return doc


@app.post("/admin/drinks")
def create_drink(drink: Drink, admin: dict = Depends(get_current_admin)):
    drink_id = create_document("drink", drink)
    return {"_id": drink_id}


@app.put("/admin/drinks/{drink_id}")
def update_drink(drink_id: str, drink: Drink, admin: dict = Depends(get_current_admin)):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    try:
        db["drink"].update_one({"_id": ObjectId(drink_id)}, {"$set": drink.model_dump()})
    except Exception:
        raise HTTPException(status_code=404, detail="Drink not found")
    return {"status": "ok"}


@app.delete("/admin/drinks/{drink_id}")
def delete_drink(drink_id: str, admin: dict = Depends(get_current_admin)):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    db["drink"].delete_one({"_id": ObjectId(drink_id)})
    return {"status": "ok"}


# Theory endpoints
@app.get("/theory")
def list_theory():
    if db is None:
        return []
    items = list(db["theorysection"].find({}).sort("order"))
    for it in items:
        it["_id"] = str(it["_id"])  # jsonify
    return items


@app.post("/admin/theory")
def create_theory(section: TheorySection, admin: dict = Depends(get_current_admin)):
    sid = create_document("theorysection", section)
    return {"_id": sid}


@app.put("/admin/theory/{section_id}")
def update_theory(section_id: str, section: TheorySection, admin: dict = Depends(get_current_admin)):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    db["theorysection"].update_one({"_id": ObjectId(section_id)}, {"$set": section.model_dump()})
    return {"status": "ok"}


@app.delete("/admin/theory/{section_id}")
def delete_theory(section_id: str, admin: dict = Depends(get_current_admin)):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    db["theorysection"].delete_one({"_id": ObjectId(section_id)})
    return {"status": "ok"}


# Certificates endpoints
@app.get("/certificates")
def list_certificates(kind: Optional[str] = None):
    if db is None:
        return []
    query = {"revoked": {"$ne": True}}
    if kind in ("course", "master"):
        query["kind"] = kind
    items = list(db["certificate"].find(query).sort("date", -1))
    for it in items:
        it["_id"] = str(it["_id"])  # jsonify
    return items


@app.post("/admin/certificates")
def create_certificate(cert: Certificate, admin: dict = Depends(get_current_admin)):
    cid = create_document("certificate", cert)
    return {"_id": cid}


@app.put("/admin/certificates/{cert_id}")
def update_certificate(cert_id: str, cert: Certificate, admin: dict = Depends(get_current_admin)):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    db["certificate"].update_one({"_id": ObjectId(cert_id)}, {"$set": cert.model_dump()})
    return {"status": "ok"}


@app.delete("/admin/certificates/{cert_id}")
def delete_certificate(cert_id: str, admin: dict = Depends(get_current_admin)):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    db["certificate"].delete_one({"_id": ObjectId(cert_id)})
    return {"status": "ok"}


# Shareable links for certificates
import secrets

class ShareCreateRequest(BaseModel):
    certificate_ids: List[str]
    expires_minutes: Optional[int] = None
    one_time: bool = False


@app.post("/admin/certificates/share")
def create_share_link(payload: ShareCreateRequest, admin: dict = Depends(get_current_admin)):
    token = secrets.token_urlsafe(24)
    expires_at = None
    if payload.expires_minutes:
        expires_at = (datetime.now(timezone.utc) + timedelta(minutes=payload.expires_minutes)).isoformat()
    link = ShareLink(token=token, certificate_ids=payload.certificate_ids, expires_at=expires_at, one_time=payload.one_time)
    _id = create_document("sharelink", link)
    return {"token": token, "id": _id}


@app.get("/share/{token}")
def public_share(token: str):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    link = db["sharelink"].find_one({"token": token})
    if not link or link.get("revoked"):
        raise HTTPException(status_code=404, detail="Link not found")
    if link.get("expires_at") and datetime.fromisoformat(link["expires_at"]) < datetime.now(timezone.utc):
        raise HTTPException(status_code=410, detail="Link expired")
    if link.get("one_time") and link.get("used"):
        raise HTTPException(status_code=410, detail="Link already used")
    ids = [ObjectId(i) for i in link.get("certificate_ids", [])]
    certs = list(db["certificate"].find({"_id": {"$in": ids}}))
    for c in certs:
        c["_id"] = str(c["_id"])  # jsonify
    if link.get("one_time"):
        db["sharelink"].update_one({"_id": link["_id"]}, {"$set": {"used": True}})
    return {"certificates": certs}


@app.post("/admin/share/{token}/revoke")
def revoke_share(token: str, admin: dict = Depends(get_current_admin)):
    if db is None:
        raise HTTPException(status_code=500, detail="Database not configured")
    db["sharelink"].update_one({"token": token}, {"$set": {"revoked": True}})
    return {"status": "revoked"}


# QR code generation for share links
@app.get("/admin/share/{token}/qrcode")
def generate_qr(token: str, admin: dict = Depends(get_current_admin)):
    base_url = os.getenv("PUBLIC_FRONTEND_URL") or os.getenv("VITE_FRONTEND_URL") or ""
    url = f"{base_url}/share/{token}" if base_url else token
    img = qrcode.make(url)
    buf = BytesIO()
    img.save(buf, format="PNG")
    b64 = base64.b64encode(buf.getvalue()).decode('ascii')
    return {"data_url": "data:image/png;base64," + b64}


# File uploads for certificate images (simple local storage)
@app.post("/admin/upload")
def upload_file(file: UploadFile = File(...), admin: dict = Depends(get_current_admin)):
    ext = os.path.splitext(file.filename)[1]
    name = f"cert_{datetime.now(timezone.utc).timestamp()}".replace(".", "_") + ext
    path = os.path.join(UPLOAD_DIR, name)
    with open(path, "wb") as f:
        f.write(file.file.read())
    return {"url": f"/static/{name}"}


# Minimal importer for Word document
from docx import Document

@app.post("/admin/import-docx")
def import_docx(doc_path: str, admin: dict = Depends(get_current_admin)):
    if not os.path.exists(doc_path):
        raise HTTPException(status_code=404, detail="Document not found")
    document = Document(doc_path)
    content = []
    for p in document.paragraphs:
        text = p.text.strip()
        if text:
            content.append(text)
    section = TheorySection(
        slug="imported",
        title_it="Teoria Importata",
        title_en="Imported Theory",
        content_it="\n\n".join(content),
        content_en="\n\n".join(content),
        order=0,
    )
    sid = create_document("theorysection", section)
    return {"theory_section_id": sid, "paragraphs": len(content)}


@app.post("/admin/import-docx-upload")
def import_docx_upload(file: UploadFile = File(...), admin: dict = Depends(get_current_admin)):
    # Read file stream directly
    document = Document(file.file)
    content = []
    for p in document.paragraphs:
        text = p.text.strip()
        if text:
            content.append(text)
    section = TheorySection(
        slug=f"imported-{int(datetime.now(timezone.utc).timestamp())}",
        title_it="Teoria Importata",
        title_en="Imported Theory",
        content_it="\n\n".join(content),
        content_en="\n\n".join(content),
        order=0,
    )
    sid = create_document("theorysection", section)
    return {"theory_section_id": sid, "paragraphs": len(content)}


# Schema endpoint for tooling
@app.get("/schema")
def get_schema():
    return {
        "collections": [
            "user",
            "drink",
            "theorysection",
            "certificate",
            "sharelink",
        ]
    }


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
